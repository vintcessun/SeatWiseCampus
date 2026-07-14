# -*- coding: utf-8 -*-
"""重构《SeatWise 进度汇报》为纯文字版：去掉全部截图，用文字讲清当前设计。
输出仓库根目录 SeatWise_进度汇报.docx（并由外部转存 .doc）。
"""
import os
from datetime import date
import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

REPO = r"D:\Scripts\SeatWiseCampus"
OUT_DOCX = os.path.join(REPO, "SeatWise_进度汇报.docx")
FONT = "微软雅黑"


def _ea(run, font):
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts"); rPr.append(rFonts)
    for a in ("w:eastAsia", "w:ascii", "w:hAnsi"):
        rFonts.set(qn(a), font)


def run_fmt(run, size=11, bold=False, color=None, italic=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    _ea(run, FONT)


doc = docx.Document()
doc.styles["Normal"].font.name = FONT
doc.styles["Normal"].font.size = Pt(11)


def title(text, size, color=None, before=0, after=6, align=WD_ALIGN_PARAGRAPH.LEFT, bold=True):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    run_fmt(p.add_run(text), size=size, bold=bold, color=color)
    return p


def h1(text):
    p = title(text, 15, color="1F3864", before=14, after=6)
    p.paragraph_format.line_spacing = 1.3
    # 底部细线
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "2"); bottom.set(qn("w:color"), "3B6CFF")
    pbdr.append(bottom); pPr.append(pbdr)
    return p


def h2(text):
    return title(text, 12.5, color="2E5496", before=8, after=3)


def para(text, indent=True):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.first_line_indent = Pt(22)
    run_fmt(p.add_run(text), size=11)
    return p


def bullet(label, text=""):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing = 1.4
    p.paragraph_format.space_after = Pt(2)
    if label:
        run_fmt(p.add_run(label + ("：" if text else "")), size=11, bold=True, color="1F3864")
    if text:
        run_fmt(p.add_run(text), size=11)
    return p


# ==================== 封面 ====================
title("SeatWise Campus", 30, color="3B6CFF", before=90, after=8, align=WD_ALIGN_PARAGRAPH.CENTER)
title("智能校园自习室预约管理平台", 20, before=0, after=6, align=WD_ALIGN_PARAGRAPH.CENTER)
title("阶段进度汇报 · 设计说明（纯文字版）", 14, color="666666", before=0, after=40,
      align=WD_ALIGN_PARAGRAPH.CENTER, bold=False)
_p = doc.add_paragraph(); _p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_fmt(_p.add_run(f"汇报日期：{date.today().isoformat()}　｜　版本：MVP / MVP+ 可运行版"), size=11, color="444444")
doc.add_page_break()


# ==================== 一、项目概述 ====================
h1("一、项目概述")
para("SeatWise Campus 是一套面向高校多校区、多楼栋、多楼层场景的 C/S 架构自习室在线预约系统，"
     "旨在解决“空位不透明、线下占座、到场无座、使用率难统计、状态更新滞后、爽约不公平”等痛点。"
     "系统提供在线选座、限时预约、到场签到、超时释放、爽约黑名单、实时热力看板与统计报表，"
     "并在此基础上实现积分排行、附近空位推荐、候补补位、组队预约、AI 选座助手，以及一套跨平台原生移动端 App（Android / iOS）等增强能力。")
para("系统的设计围绕两个技术核心展开：一是“时间片并发选座的正确性”，二是“多客户端实时看板的一致性”。"
     "所有其他功能均服务于这两条主线，且始终坚持“后端权威、MySQL 为唯一正确性来源”的原则。")


# ==================== 二、系统总体架构 ====================
h1("二、系统总体架构与技术栈")
para("系统采用经典的 C/S（浏览器/服务器）分层架构。前端只通过 REST 接口与 SSE 通道访问后端，"
     "不直接接触数据库；后端集中承载业务规则、并发控制、权限鉴权与数据一致性。"
     "存储层职责严格分离：MySQL 8 是最终正确性来源，Redis 7 / Redisson 仅承担缓存、计数、分布式锁与调度，二者不可混淆。")
h2("2.1  技术栈")
bullet("后端", "JDK 21 + Spring Boot 3.3.5 + MyBatis-Plus 3.5.9 + MySQL 8 + Redis 7 + "
               "Redisson 3.37.0 + Sa-Token 1.39.0 + Knife4j 4.5.0。")
bullet("前端", "Vue 3 + Vite + Element Plus + ECharts + Pinia + Vue Router + Axios；"
               "自研轻量国际化（中/英）与无障碍层。")
bullet("部署", "Docker Compose 一键启动，含 MySQL、Redis、后端、前端 Nginx 四个容器，附健康检查与自动初始化 SQL。")
h2("2.2  分层与边界")
bullet("前端 client", "页面层 → 组件层 → Pinia 状态管理 → Axios REST / SSE 连接层；只做交互与表单校验，最终结果以后端返回为准。")
bullet("后端 server", "Controller（REST + SSE 端点）→ Sa-Token 鉴权 → Service 业务层 → MyBatis-Plus Mapper；"
                     "另有 Redisson 锁/延迟能力、定时任务与 SSE 推送管理器。")
bullet("通信", "REST（HTTP/JSON，统一响应 code/message/data/traceId）用于请求-响应；SSE（text/event-stream）用于服务端→客户端的座位状态增量推送。")


# ==================== 三、数据模型与存储 ====================
h1("三、数据模型与存储设计")
para("基础数据以“校区 campus → 楼栋 building → 自习室 study_room → 座位 seat”四级层次组织；"
     "座位以行列网格建模，cell_type 区分 SEAT/AISLE/EMPTY/DISABLED，并用 enabled 表示运营态启停。"
     "预约业务由“预约主记录 reservation + 时间片占用 reservation_slot”两张表承载，另辅以黑名单、积分流水、日聚合、通知等表。")
h2("3.1  时间片拆片与唯一索引兜底（关键）")
para("一次预约按 30 分钟粒度被拆分为若干时间片，逐片写入 reservation_slot 表。"
     "该表在 (seat_id, date, slot_index) 上建立唯一索引 uk_seat_date_slot，是防止重复预约的最终兜底："
     "同座同片的并发插入由数据库原子判重，冲突者被唯一约束直接拒绝。释放预约即删除对应时间片行。")
para("之所以不采用“读取时间区间再比较重叠”的方式，是因为区间重叠判断需先读后比较，并发下存在检查—写入竞态，"
     "且难以施加数据库级唯一约束；而拆片后“同座同片”天然可用唯一索引，插入即判重，正确性由数据库保证。")


# ==================== 四、核心设计一：并发选座 ====================
h1("四、核心设计一：时间片并发选座的正确性")
para("预约提交在服务端完成如下流程：先校验登录、角色、黑名单、开放时间、单日次数与用户自身时段冲突；"
     "将起止时间换算为时间片序号列表；对目标座位用 Redisson 分布式锁降低冲突；"
     "随后在同一数据库事务内插入预约主记录并批量插入时间片占用记录。")
bullet("成功", "提交事务，写 Redis 座位状态=RESERVED，登记超时释放，推送 SSE seat_reserved 事件，返回预约成功。")
bullet("冲突", "命中唯一索引即回滚事务，返回 SEAT_ALREADY_RESERVED。")
para("这里的关键取舍是：Redisson 锁只负责“减少冲突、降低失败率”，而 MySQL 唯一索引才是“最终兜底”。"
     "即使 Redis 不可用或锁失效，数据库唯一约束仍能保证同座同片至多一人成功，绝不双占。")
h2("4.1  预约状态机")
para("一条预约的生命周期由后端状态机驱动：预约成功进入“待签到 PENDING_SIGN_IN”；在签到窗口内签到转为“使用中 IN_USE”；"
     "主动签退或到达结束时间自动完成转为“已完成 COMPLETED”；待签到期间取消转为“已取消 CANCELLED”；"
     "超时未签到由定时任务置为“爽约释放 EXPIRED_RELEASED”。所有状态变更均在后端事务内完成并触发 SSE 推送，"
     "座位不会卡死在“使用中”。")


# ==================== 五、核心设计二：实时一致性 ====================
h1("五、核心设计二：多客户端实时看板一致性")
para("实时热力看板采用“初始化快照 + SSE 增量推送”方案。客户端打开看板时先拉取一次 board_snapshot"
     "（座位网格及每格状态），随后建立 SSE 订阅；此后任一客户端完成预约/签到/释放等操作，"
     "服务端生成对应增量事件（seat_reserved / seat_in_use / seat_released 等）并广播给所有订阅者，"
     "各端仅对相关座位做局部更新，实现多客户端秒级一致，避免整屏轮询。")
bullet("断线重连", "连接中断后前端自动重连并重新拉取快照，保证补齐期间遗漏的状态。")
bullet("心跳保活", "服务端每 15 秒向所有 SSE 连接发送心跳，及时清理失效连接。")


# ==================== 六、座位生命周期治理 ====================
h1("六、座位生命周期治理：签到 / 超时 / 黑名单")
para("为杜绝“占而不用”，系统对座位生命周期做闭环治理。当前实现以短周期定时扫描为主（文档既定的兜底方案，演示更直观可控）："
     "每 5 秒扫描一次，处理超时未签到的自动释放、使用中到点的自动完成、候补名额过期回收与到时提醒。")
bullet("签到窗口", "预约开始后若在签到窗口内（默认约 15 分钟）未签到，则自动释放该预约并置座位为 FREE、推送 SSE。")
bullet("爽约与黑名单", "超时未签到累计爽约计数并扣分；达到阈值自动写入 blacklist_record，使该用户在限制期（7 天）内预约被拒，但仍可登录与查看历史。")
bullet("自动完成", "使用中的预约到达结束时间后由系统自动完成并释放座位，避免长期占用。")


# ==================== 七、已实现功能模块 ====================
h1("七、已实现功能模块总览")
h2("7.1  MVP（P1–P6）")
bullet("账号与鉴权", "Sa-Token 登录、注册、找回密码、图形验证码；STUDENT / ADMIN 角色权限矩阵。")
bullet("基础数据与座位排布", "校区/楼栋/自习室/开放时间维护；座位行列网格与启用禁用、可视化排布编辑。")
bullet("核心预约", "30 分钟时间片、并发控制（Redisson 锁 + 唯一索引兜底）、单日次数限制。")
bullet("座位管控", "签到、签退、主动取消、超时自动释放、结束自动完成。")
bullet("黑名单", "爽约达阈值自动进入黑名单并限制预约。")
bullet("实时看板", "快照 + SSE 增量，多客户端秒级同步。")
bullet("数据报表", "预约状态分布、热门时段、取消率/爽约率、利用率排行（ECharts）。")
h2("7.2  MVP+（P7–P8）与扩展（P9 部分落地）")
bullet("积分排行", "守约加分、爽约扣分，生成积分排行榜。")
bullet("附近空位推荐", "按“同楼栋 > 距离最近 > 空位更多”推荐可用自习室。")
bullet("候补与组队", "满座候补自动补位；组队原子预约（同一事务多座位同成或同败）。")
bullet("智能与辅助", "AI 选座助手（自然语言需求 → 可解释推荐，离线规则引擎，可接入大模型）、站内通知中心、"
                   "临时占座、历史回放、时空占用图、专注番茄钟、自习报告。")


# ==================== 八、前端设计 ====================
h1("八、前端设计与体验")
bullet("状态与数据流", "Pinia 管理登录态与视图状态；座位状态一律来自“快照 + SSE”，前端不自行推算占用。")
bullet("国际化", "自研轻量 i18n，支持中/英一键切换并联动 Element Plus 组件与页面 lang 属性。")
bullet("无障碍", "跳转主内容 skip-link、键盘可见焦点环、图标按钮 aria 标签、尊重系统减少动态效果偏好。")
bullet("移动端响应式", "窄屏下侧边栏收起为抽屉式导航、栅格纵向堆叠，桌面/iOS/Android 三端均无横向溢出。")
bullet("主题", "支持明亮 / 深色主题切换。")


# ==================== 九、跨平台移动端 App ====================
h1("九、跨平台移动端 App（Android / iOS）")
para("除响应式 Web 外，项目另用 Kotlin Multiplatform（KMP）+ Compose Multiplatform（CMP）实现了一套原生移动端 App，"
     "做到“一套代码同时构建 Android（APK）与 iOS（IPA）”，复用 Web 后端的 REST 接口，认证同样复用 satoken 头，"
     "从而在原有 C/S 架构上补齐了原生手机端入口。")
h2("9.1  技术与结构")
bullet("技术栈", "Kotlin 2.0.21 + Compose Multiplatform 1.7.1 + Ktor 2.3.12（Android 用 OkHttp 引擎、iOS 用 Darwin 引擎）"
                "+ kotlinx.serialization/coroutines + Gradle 8.9 / AGP 8.5.2。")
bullet("工程结构", "shared 共享模块（数据模型 Models、Ktor 客户端 Api、Compose UI 的 App）+ androidApp（产出 APK）"
                 "+ iosApp（SwiftUI 承载 Compose，XcodeGen 生成 Xcode 工程，产出 IPA）。")
h2("9.2  已实现功能（学生端子集）")
bullet("", "登录（登录页可填写后端地址，内置演示快捷登录）、自习室列表、座位网格点击空闲座位一键预约、我的预约与取消。")
h2("9.3  构建与产物")
bullet("CI 双端构建", "仓库根 .github/workflows/mobile-build.yml：ubuntu 构建 Android APK、macOS 构建 iOS IPA，"
                    "push mobile/** 或手动触发即可运行。")
bullet("已产出二进制", "已通过 CI 产出并置于 dist-mobile/：androidApp-release.apk（CI 演示 keystore 签名，minSdk 24，"
                    "Android 7.0+ 可直接安装）、SeatWise-unsigned.ipa（未签名 IPA，真机分发需 Apple 证书或侧载）。")
bullet("待办", "移动端默认后端地址当前留空，需在后端正式部署并给出访问地址后填入并重新构建 CI，产物即内置该地址、用户无需手填。")


# ==================== 十、测试与验证 ====================
h1("十、测试与验证结果")
para("通过“自动化脚本验证核心正确性 + 真实浏览器多端页面走查”两种方式对系统做端到端验证，核心红线全部通过：")
bullet("冒烟测试（13 项）", "登录、看板、预约、重复预约拒绝、签到签退、座位释放、单日限次、报表、权限等全部通过。")
bullet("并发抢座", "8 个学生并发抢同一座位同一时段，仅 1 人成功、7 人被拒。")
bullet("实时推送", "订阅看板后触发预约，另一端秒级收到 seat_reserved 并局部变色。")
bullet("超时与黑名单", "未签到超时释放并累计爽约触发黑名单，随后预约被拒（USER_IN_BLACKLIST）。")
bullet("多端与双语", "学生端 9 页、管理端 10 页在桌面/iOS/Android、中英双语、明暗主题下走查，均无控制台错误。")


# ==================== 十一、部署与运行 ====================
h1("十一、部署与运行")
bullet("一键启动", "docker compose up -d --build。")
bullet("访问入口", "前端 http://localhost:8888；后端 API http://localhost:18080；接口文档 Knife4j http://localhost:18080/doc.html。")
bullet("演示账号", "admin / admin123（管理员），student1~8 / 123456（学生），登录页提供快捷登录。")
bullet("端口说明", "宿主机 8080 被占用，后端对外映射到 18080；浏览器仅需访问 8888，由 Nginx 反代到后端。")


# ==================== 十二、进度与取舍 ====================
h1("十二、当前进度、务实取舍与后续计划")
bullet("里程碑", "P0 文档与脚手架、P1–P6（MVP）、P7–P8（MVP+）均已完成并可运行；P9 扩展中 AI 助手、通知、候补、组队、番茄钟、回放等已落地，校园地图等仅接口预留。")
bullet("超纲交付", "原计划非目标的“移动端原生 App”已额外用 KMP + Compose Multiplatform 落地为 Android/iOS 双端并产出可安装包。")
bullet("务实取舍", "为演示稳定与生态兼容，Spring Boot 采用 3.3.5；超时释放采用短周期定时扫描（文档定义的兜底方案）；报表采用实时聚合。以上均在代码注释与文档中标注，不影响功能演示。")
bullet("后续计划", "将超时释放升级为 Redisson 延迟队列；报表改为 room_daily_stats 聚合表 + 定时任务以提升大数据量性能；补充管理端基础数据表单与积分规则配置；引入自动化测试到 CI 并扩展并发压测。")


doc.save(OUT_DOCX)
print("已生成:", OUT_DOCX)
