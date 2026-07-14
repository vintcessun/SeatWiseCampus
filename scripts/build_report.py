# -*- coding: utf-8 -*-
"""在《Java程序设计实践》项目开发报告（团队版）模板中填充全部正文：
  一、项目组合影和分工介绍（留空表格 + 合影占位，供用户填写）
  二、项目目标与要求
  三、项目内容 → （一）功能需求 / （二）系统设计 / （三）系统测试
删除模板所有占位（XXXX、"（字体为宋体…）"、示例图表、格式说明），按“图居中·图题在下、表居中·表题在上”重排。
输出两份：*_完整版.docx（正文+图+表）、*_仅插图版.docx（标题+图/表+题，供 Word with Claude 续写）。
"""
import os
from PIL import Image
import docx
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

REPO = r"D:\Scripts\SeatWiseCampus"
TEMPLATE = r"C:\Users\vintc\OneDrive - stu.xmu.edu.cn\学校\课程\大二小学期\java\3《Java程序设计实践》项目开发报告（团队版）.docx"
OUT_DIR = os.path.dirname(TEMPLATE)
OUT_FULL = os.path.join(OUT_DIR, "3《Java程序设计实践》项目开发报告（团队版）_完整版.docx")
OUT_IMG = os.path.join(OUT_DIR, "3《Java程序设计实践》项目开发报告（团队版）_仅插图版.docx")

FIG = os.path.join(REPO, "scripts", "figures")
BODY_FONT = "宋体"
HEAD_FONT = "黑体"
CONTENT_W_CM = 15.4
PORTRAIT_H_CM = 10.0

# 全局连续编号
_fig = [0]
_tbl = [0]
def next_fig(): _fig[0] += 1; return f"图 {_fig[0]}"
def next_tbl(): _tbl[0] += 1; return f"表 {_tbl[0]}"


def sp(p):
    return os.path.join(REPO, p)


# ---------- 字体 / 段落工具 ----------
def _ea(run, font):
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts"); rPr.append(rFonts)
    for a in ("w:eastAsia", "w:ascii", "w:hAnsi"):
        rFonts.set(qn(a), font)


def _fmt(run, font=BODY_FONT, size=12, bold=False, italic=False, color=None):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    _ea(run, font)


def _set_table_borders(tbl):
    tblPr = tbl._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement("w:" + edge)
        e.set(qn("w:val"), "single"); e.set(qn("w:sz"), "4")
        e.set(qn("w:space"), "0"); e.set(qn("w:color"), "8A93A6")
        borders.append(e)
    tblPr.append(borders)


def _fix_layout(tbl, widths_cm):
    """固定表格布局，使列宽严格生效（否则 Word 会按内容自动伸缩）。"""
    tbl.autofit = False
    tbl.allow_autofit = False
    tblPr = tbl._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)
    total = int(sum(widths_cm) * 567)  # cm → twips
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:type"), "dxa"); tblW.set(qn("w:w"), str(total))
    tblPr.append(tblW)
    # 网格列宽
    grid = tbl._tbl.find(qn("w:tblGrid"))
    if grid is not None:
        for gc, wcm in zip(grid.findall(qn("w:gridCol")), widths_cm):
            gc.set(qn("w:w"), str(int(wcm * 567)))


def _shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def _code_font(run):
    run.font.size = Pt(9)
    rPr = run._element.get_or_add_rPr()
    rF = rPr.find(qn("w:rFonts"))
    if rF is None:
        rF = OxmlElement("w:rFonts"); rPr.append(rF)
    rF.set(qn("w:ascii"), "Consolas")
    rF.set(qn("w:hAnsi"), "Consolas")
    rF.set(qn("w:eastAsia"), "宋体")


# ---------- 封面（项目组名 / 选题 / 成员） ----------
MEMBERS = [
    ("王子恒", "34520242201240", "总负责、代码"),
    ("林俊豪", "22920242203293", "端到端测试"),
    ("林润宸", "22920242203296", "代码"),
    ("吕宜家", "22920242203340", "前端测试、PPT 制作"),
    ("谭嘉亮", "22920242203390", "后端测试、PPT 制作"),
    ("叶雨扬", "22920242203485", "安卓端、苹果端测试"),
    ("莫岢毅", "22920242203354", "代码"),
]
GROUP_NAME = "大梦一瓜"
TOPIC = "题目二　智能校园自习室预约管理平台"
PHOTO = os.path.join(OUT_DIR, "合照.png")  # 项目组合影


def _underlined_run(par):
    """返回该填空段中带下划线的空格 run（即下划线本身）。"""
    for r in par.runs:
        if r.font.underline:
            return r
    return par.runs[-1] if par.runs else None


def _fill_line(par, value):
    """把值填到下划线上：只改下划线 run 文本，保留其下划线/加粗/字号格式。"""
    r = _underlined_run(par)
    if r is None:
        return
    orig = r.text
    tail = " " * max(6, len(orig) - len(value) - 2)  # 保留一段延伸下划线
    r.text = " " + value + tail


def _fill_member(par, name, num):
    """姓名学号行：下划线含全角空格分隔姓名区/学号区，分别填入。"""
    r = _underlined_run(par)
    if r is None:
        return
    orig = r.text
    if "　" in orig:
        left, right = orig.split("　", 1)
        lpad = " " * max(2, len(left) - len(name) - 1)
        rpad = " " * max(2, len(right) - len(num) - 1)
        r.text = " " + name + lpad + "　" + " " + num + rpad
    else:
        r.text = " " + name + "    " + num + "     "


def fill_title(doc):
    import re
    from copy import deepcopy
    topic_par = group_par = None
    slots = []
    for ch in list(doc.element.body.iterchildren()):
        if not ch.tag.endswith("}p"):
            continue
        par = Paragraph(ch, doc._body)
        squeezed = re.sub(r"\s", "", par.text)  # 去除半角与全角空白
        # 精确匹配封面填空标签（避免匹配到正文“项目组合影和分工介绍”标题）
        if squeezed == "课程项目":
            topic_par = par
        elif squeezed == "项目组":
            group_par = par
        elif squeezed == "姓名学号":
            slots.append(par)
    # 成员多于模板槽位时，按最后一个空槽复制补足（保留下划线格式）
    if slots and len(MEMBERS) > len(slots):
        last_el = slots[-1]._p
        for _ in range(len(MEMBERS) - len(slots)):
            newp = deepcopy(last_el)
            last_el.addnext(newp)
            last_el = newp
            slots.append(Paragraph(newp, doc._body))
    if topic_par:
        _fill_line(topic_par, TOPIC)
    if group_par:
        _fill_line(group_par, GROUP_NAME)
    for i, par in enumerate(slots):
        if i < len(MEMBERS):
            _fill_member(par, MEMBERS[i][0], MEMBERS[i][1])
        else:
            _fill_member(par, "", "")


class Cursor:
    def __init__(self, doc, anchor_p, images_only):
        self.doc = doc
        self.cur = anchor_p
        self.images_only = images_only

    def _add_p_after(self):
        p = OxmlElement("w:p")
        self.cur.addnext(p)
        self.cur = p
        return Paragraph(p, self.doc._body)

    # 小节标题（如 2.1）
    def sub(self, text, size=13):
        para = self._add_p_after()
        para.paragraph_format.space_before = Pt(10)
        para.paragraph_format.space_after = Pt(4)
        para.paragraph_format.line_spacing = 1.4
        _fmt(para.add_run(text), font=HEAD_FONT, size=size, bold=True, color="000000")
        return para

    # 正文（images_only 时跳过）
    def body(self, text):
        if self.images_only:
            return None
        para = self._add_p_after()
        para.paragraph_format.line_spacing = 1.5
        para.paragraph_format.space_after = Pt(4)
        para.paragraph_format.first_line_indent = Pt(24)
        _fmt(para.add_run(text), font=BODY_FONT, size=12, color="000000")
        return para

    def bullet(self, label, text=""):
        if self.images_only:
            return None
        para = self._add_p_after()
        para.paragraph_format.line_spacing = 1.45
        para.paragraph_format.left_indent = Pt(24)
        r = para.add_run("· ")
        _fmt(r, font=BODY_FONT, size=12, bold=True, color="3B6CFF")
        if label:
            _fmt(para.add_run(label + ("：" if text else "")), size=12, bold=True, color="000000")
        if text:
            _fmt(para.add_run(text), size=12, color="000000")
        return para

    def image(self, filename, max_w_cm=None):
        path = filename if os.path.isabs(filename) else os.path.join(FIG, filename)
        para = self._add_p_after()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(2)
        run = para.add_run()
        if not os.path.exists(path):
            _fmt(para.add_run(f"（图缺失：{os.path.basename(path)}）"), color="FF0000")
            return para
        w, h = Image.open(path).size
        if h > w * 1.25:
            run.add_picture(path, height=Cm(PORTRAIT_H_CM))
        else:
            run.add_picture(path, width=Cm(max_w_cm or CONTENT_W_CM))
        return para

    def caption(self, text):
        para = self._add_p_after()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_after = Pt(10)
        _fmt(para.add_run(text), font=BODY_FONT, size=10.5, color="595959")
        return para

    def figure(self, imgfile, name, max_w_cm=None):
        self.image(imgfile, max_w_cm=max_w_cm)
        self.caption(f"{next_fig()}    {name}")

    def code(self, label, text):
        """核心代码块（带浅底纹与边框的等宽代码）。"""
        if label:
            lab = self._add_p_after()
            lab.paragraph_format.space_before = Pt(4)
            lab.paragraph_format.space_after = Pt(2)
            _fmt(lab.add_run("核心代码：" + label), font=BODY_FONT, size=10.5, bold=True, color="000000")
        para = self._add_p_after()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.left_indent = Pt(10)
        para.paragraph_format.space_after = Pt(8)
        para.paragraph_format.line_spacing = 1.05
        pPr = para._p.get_or_add_pPr()
        shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), "F5F6FA")
        pPr.append(shd)
        pbdr = OxmlElement("w:pBdr")
        for edge in ("top", "left", "bottom", "right"):
            e = OxmlElement("w:" + edge)
            e.set(qn("w:val"), "single"); e.set(qn("w:sz"), "4")
            e.set(qn("w:space"), "4"); e.set(qn("w:color"), "D0D5DD")
            pbdr.append(e)
        pPr.append(pbdr)
        lines = text.strip("\n").split("\n")
        for i, ln in enumerate(lines):
            r = para.add_run(ln if ln else " ")
            _code_font(r)
            if i < len(lines) - 1:
                r.add_break()
        return para

    def table(self, header, rows, title, col_w=None):
        cap = self._add_p_after()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_before = Pt(6)
        _fmt(cap.add_run(f"{next_tbl()}    {title}"), font=BODY_FONT, size=10.5, bold=True, color="595959")
        tbl = self.doc.add_table(rows=1, cols=len(header))
        _set_table_borders(tbl)
        tbl.alignment = 1
        for j, htext in enumerate(header):
            c = tbl.rows[0].cells[j]
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            _fmt(c.paragraphs[0].add_run(htext), font=HEAD_FONT, size=10.5, bold=True, color="000000")
            _shade(c, "D9E2F3")
        for row in rows:
            cells = tbl.add_row().cells
            for j, val in enumerate(row):
                _fmt(cells[j].paragraphs[0].add_run(str(val)), font=BODY_FONT, size=10.5, color="000000")
        if col_w:
            for j, wcm in enumerate(col_w):
                for r in tbl.rows:
                    r.cells[j].width = Cm(wcm)
            _fix_layout(tbl, col_w)
        self.cur.addnext(tbl._tbl)
        self.cur = tbl._tbl
        self._add_p_after()
        return tbl


# ---------- 定位与清理 ----------
def find_p(doc, text):
    for ch in doc.element.body.iterchildren():
        if ch.tag.endswith("}p"):
            if Paragraph(ch, doc._body).text.strip() == text:
                return ch
    return None


def delete_between(start_el, end_el):
    node = start_el.getnext()
    while node is not None and node is not end_el:
        nxt = node.getnext()
        node.getparent().remove(node)
        node = nxt


def delete_after_until_sectpr(start_el):
    node = start_el.getnext()
    while node is not None:
        nxt = node.getnext()
        if node.tag.endswith("}sectPr"):
            break
        node.getparent().remove(node)
        node = nxt


# ============================================================ 各部分内容
def build_photo(c):
    c.body(f"本项目为《Java 程序设计实践》课程项目，项目组名称为“{GROUP_NAME}”，选题为“{TOPIC}”。"
           f"项目由项目组 {len(MEMBERS)} 名成员共同完成，成员及主要分工如下表所示。")
    c.table(
        ["姓名", "学号", "主要分工"],
        [[m[0], m[1], m[2]] for m in MEMBERS],
        "项目成员与分工",
        col_w=[3.0, 5.0, 7.4],
    )
    c.body("项目组全体成员合影如下图所示。")
    c.figure(PHOTO, "项目组合影", max_w_cm=12.0)


def build_goal(c):
    c.body(f"本项目选题为“{TOPIC}”，旨在构建一套面向高校多校区、多楼栋、多楼层场景的 C/S 架构自习室在线预约系统 —— SeatWise Campus，"
           "解决“空位不透明、线下占座、到场无座、使用率难统计、状态更新滞后、爽约不公平”等痛点，"
           "让学生能方便地找到并锁定空位、让管理员掌握使用情况、用规则保障使用公平。")
    c.sub("（一）功能目标", 12.5)
    c.bullet("学生侧", "在线筛选自习室、按 30 分钟时间片选座预约、到场签到、签退、取消，并查看本人预约、积分与附近空位。")
    c.bullet("管理侧", "维护校区/楼栋/自习室与座位排布，查看实时座位看板与统计报表，管理黑名单。")
    c.bullet("公平保障", "超时未签到自动释放、爽约计数与黑名单、积分激励，减少占而不用。")
    c.sub("（二）非功能要求", 12.5)
    c.bullet("正确性", "并发选座下同一座位同一时间片绝不双占（时间片唯一索引最终兜底）。")
    c.bullet("实时性", "座位状态变更后其他客户端秒级可见（初始化快照 + SSE 增量推送）。")
    c.bullet("性能", "报表读聚合口径、看板读缓存+快照，避免全表实时扫描。")
    c.bullet("可用性与安全", "SSE 断线自动重连；依赖故障不导致双占；Sa-Token 登录态与角色鉴权。")
    c.bullet("可部署性", "Docker Compose 一键启动 MySQL / Redis / 后端 / 前端四容器。")


def build_content_intro(c):
    c.body("本章从功能需求、系统设计与系统测试三个方面介绍项目内容。系统的两大技术核心是"
           "“时间片并发选座的正确性”与“多客户端实时看板的一致性”，其余功能均围绕二者展开。")


def build_func(c):
    c.body("依据需求分析，系统功能按模块划分如下（分级：MVP 为最小可用版本，MVP+ 为增强，扩展为后续能力），核心功能一览见下表。")
    c.table(
        ["功能模块", "主要功能", "分级"],
        [
            ["自习室管理", "校区/楼栋/楼层/自习室与开放时间维护；座位行列网格排布（SEAT/AISLE/EMPTY/DISABLED）与启用禁用", "MVP"],
            ["学生预约", "登录；按校区/楼栋/楼层筛选；查看空位；选日期+30 分钟时间片+座位；提交预约；单日限次；黑名单拦截；查看本人记录", "MVP"],
            ["座位管控", "待签到→签到→使用中；主动签退或到点自动完成；超时未签到自动释放；主动取消；爽约计数与黑名单", "MVP"],
            ["实时看板", "座位热力图（空闲/已约/使用中/禁用 + 本人高亮）；初始化快照 + SSE 增量，多客户端秒级同步", "MVP"],
            ["数据报表", "使用率、热门时段、取消率、爽约率、利用率排行，可按校区/楼栋/自习室筛选", "MVP"],
            ["积分排名", "守约加分、爽约扣分，按周/月生成排行榜（与黑名单解耦）", "MVP+"],
            ["位置与附近空位", "管理端接入地图选点维护楼栋经纬度；学生端浏览器定位识别最近楼栋，按“同楼栋 > 距离最近 > 空位更多”推荐可用自习室", "MVP+"],
            ["智能与扩展", "AI 选座助手、站内通知中心、满座候补、组队原子预约、专注番茄钟、自习报告、历史回放", "扩展"],
        ],
        "功能需求一览",
        col_w=[2.6, 10.4, 2.4],
    )
    c.sub("（一）自习室与座位管理（管理员）", 12.5)
    c.bullet("", "录入并维护校区、楼栋、楼层、自习室及其开放时间；以行列网格可视化编辑座位排布，区分座位/过道/空位/禁用格，并可对真实座位做运营态启用与禁用。")
    c.sub("（二）学生选座预约", 12.5)
    c.bullet("", "登录后按校区→楼栋→楼层逐级筛选自习室，查看座位空位；选择日期与预约起止时间（按 30 分钟拆片），在座位网格中自选空闲座位提交预约；系统限制单日预约次数，并拒绝黑名单用户预约。")
    c.sub("（三）座位管控与公平保障", 12.5)
    c.bullet("", "预约成功后进入待签到；在签到窗口（默认 15 分钟）内签到转为使用中，主动签退或到达结束时间自动完成并释放座位；超时未签到自动释放并计爽约，爽约达阈值进入黑名单，限制期内不能预约但仍可登录与查看历史。")
    c.sub("（四）实时看板、报表与增强", 12.5)
    c.bullet("", "实时热力看板以“快照 + SSE 增量”多端同步座位状态；数据报表统计使用率、热门时段、取消率、爽约率与利用率排行；此外提供积分排行、附近空位推荐、AI 选座助手、候补、组队、通知等增强功能。")


DESIGN_FIGS = [
    ("2.1  系统总体架构",
     "系统采用经典的 C/S（浏览器/服务器）分层架构。前端 client 工程基于 Vue3 + Vite + Element Plus + Pinia + ECharts，"
     "分为页面层、组件层、状态管理层与 Axios REST / SSE 连接层，仅通过 REST 接口与 SSE 通道访问后端，不直接接触数据库。"
     "后端 server 工程基于 Spring Boot + MyBatis-Plus + Sa-Token，由 Controller、Service、Mapper 三层构成，"
     "并配合 Redisson 分布式锁/延迟能力与定时任务。MySQL 8 作为唯一的正确性来源，Redis 7 仅承担座位状态缓存、"
     "计数与调度，二者职责严格分离。系统总体架构如图所示。",
     "fig1_architecture.png", "系统总体架构图"),
    ("2.2  数据库设计",
     "数据库以“校区—楼栋—自习室—座位”的层级组织基础数据，以“预约主记录 reservation + 时间片占用 reservation_slot”"
     "承载预约业务，并辅以黑名单、积分流水、日聚合等表。核心表关系如下图所示。其中 reservation_slot 表将一次预约按"
     "30 分钟粒度拆分为若干时间片，其 (seat_id, date, slot_index) 唯一索引是防止重复预约的最终兜底，该表字段设计见下表。",
     "fig2_er.png", "数据库 E-R 关系图"),
    ("2.3  时间片并发模型",
     "为在高并发选座下保证“同一座位同一时段只能被一人预约”，系统不采用“读取区间再比较重叠”的方式（存在检查—写入竞态、"
     "难以施加数据库级唯一约束），而是将预约按 30 分钟粒度拆片存入 reservation_slot 表。如下图所示，每个 (座位, 时间片) "
     "至多被一条预约占用，插入即判重，冲突由数据库唯一索引原子拒绝，天然规避并发双占。",
     "fig3_slot_model.png", "时间片占用模型（30 分钟粒度拆片）"),
    ("2.4  核心预约与并发控制",
     "预约提交流程如下图所示：服务端先完成登录、角色、黑名单、开放时间、单日次数及自身时段冲突等校验，将起止时间换算为"
     "时间片序号列表；随后使用 Redisson 对目标座位加分布式锁以降低冲突概率，在同一事务内插入预约主记录与批量时间片占用记录。"
     "若命中 (seat_id, date, slot_index) 唯一索引冲突则回滚并返回 SEAT_ALREADY_RESERVED；成功则提交事务、写入 Redis 座位状态、"
     "登记超时释放并推送 SSE 事件。Redisson 锁只是“减少冲突”，MySQL 唯一索引才是“最终兜底”——即使 Redis 不可用，"
     "并发预约仍不会双占。",
     "fig4_reservation_flow.png", "预约提交与并发控制流程"),
    ("2.5  预约状态机",
     "一条预约的完整生命周期由后端状态机驱动，如下图所示。预约成功后进入“待签到 PENDING_SIGN_IN”；在签到窗口内签到"
     "转为“使用中 IN_USE”，主动签退或到达结束时间自动完成转为“已完成 COMPLETED”；待签到期间用户可取消转为“已取消 "
     "CANCELLED”；若超时未签到则由定时任务置为“爽约释放 EXPIRED_RELEASED”，累计爽约并可能进入黑名单。所有状态变更均在"
     "后端事务内完成并触发 SSE 推送。",
     "fig5_state_machine.png", "预约状态机"),
    ("2.6  实时看板（快照 + SSE 增量）",
     "实时热力看板采用“初始化快照 + SSE 增量推送”的方案，如下图所示。客户端打开看板时先拉取一次 board_snapshot（座位网格"
     "及每格状态），随后建立 SSE 订阅；当任一客户端完成预约/签到/释放等操作，服务端生成对应增量事件（如 seat_reserved）并由 SSE "
     "推送管理器广播给所有订阅者，各端仅对相关座位做局部更新，实现多客户端秒级一致，避免整屏轮询。断线后前端自动重连并重取快照。",
     "fig6_sse.png", "实时看板 SSE 推送机制"),
    ("2.7  超时释放与黑名单",
     "为避免“占而不用”，系统对超时未签到的预约做自动释放。如下图所示，当预约开始后签到窗口（默认 15 分钟）内仍未签到，"
     "则释放时间片、置座位为 FREE 并推送 SSE，同时爽约计数加一、扣除积分；当爽约次数达到阈值时写入 blacklist_record，"
     "使该用户在限制期（7 天）内预约被拒，但仍可登录与查看历史。当前实现以短周期定时扫描为主（文档既定的兜底方案，演示更直观可控）。",
     "fig7_timeout.png", "超时释放与黑名单机制"),
    ("2.8  部署架构",
     "系统通过 Docker Compose 一键部署，如下图所示，包含 frontend（Nginx 提供静态资源并反向代理 /api）、backend（Spring Boot）、"
     "mysql、redis 四个容器。浏览器/移动端仅访问前端 8888 端口，Nginx 在容器网络内将接口请求反代至后端；后端分别通过 JDBC 与 "
     "Lettuce/Redisson 连接 MySQL 与 Redis。该拓扑使前后端、数据库、缓存解耦，便于本地演示与迁移。",
     "fig8_deploy.png", "部署拓扑图"),
]

DB_TABLE = (
    ["字段名称", "数据类型", "允许空", "键 / 约束", "说明"],
    [
        ["id", "BIGINT", "否", "PK 自增", "主键"],
        ["reservation_id", "BIGINT", "否", "IDX 外键", "关联的预约主记录"],
        ["seat_id", "BIGINT", "否", "UNIQUE(seat_id,date,slot_index)", "占用的座位；唯一键兜底防重"],
        ["date", "DATE", "否", "↑ 联合唯一", "预约日期"],
        ["slot_index", "INT", "否", "↑ 联合唯一", "30 分钟时间片序号"],
        ["created_time", "DATETIME", "是", "", "创建时间"],
    ],
    "reservation_slot 时间片占用表结构",
)

API_TABLE = (
    ["接口", "方法", "路径", "说明"],
    [
        ["登录", "POST", "/api/auth/login", "用户名密码登录，返回 token 与角色"],
        ["自习室查询", "GET", "/api/study-rooms", "按校区/楼栋/楼层筛选自习室"],
        ["座位排布", "PUT", "/api/study-rooms/{id}/layout", "保存座位行列网格排布"],
        ["看板快照", "GET", "/api/study-rooms/{id}/board", "返回座位网格及每格状态"],
        ["实时订阅", "GET(SSE)", "/api/board/stream", "text/event-stream 增量推送"],
        ["提交预约", "POST", "/api/reservations", "并发控制下创建预约"],
        ["签到 / 签退", "POST", "/api/reservations/{id}/check-in", "签到；签退为 .../check-out，释放座位"],
        ["取消预约", "POST", "/api/reservations/{id}/cancel", "取消并释放时间片"],
        ["我的预约", "GET", "/api/reservations/me", "本人预约记录"],
        ["数据报表", "GET", "/api/reports/summary", "使用率/取消率/爽约率等聚合"],
        ["积分排行", "GET", "/api/scores/ranking", "按周/月排行榜"],
        ["附近空位", "GET", "/api/rooms/nearest-available", "按距离+空位+开放状态推荐"],
    ],
    "核心 REST 接口一览",
)

TEST_SECTIONS = [
    ("3.1  测试环境与方法",
     "测试在 Docker Compose 环境下进行（前端 Nginx:8888、后端 Spring Boot:18080、MySQL 8、Redis 7），"
     "采用“自动化冒烟/并发脚本 + 真实 Chrome 多端页面走查”相结合的方式。自动化脚本覆盖登录、看板、预约、重复预约拒绝、"
     "8 并发仅 1 成功、签到签退、座位释放、单日限次、报表、权限等 13 项核心用例并全部通过，地图选点、定位推荐与专注番茄钟"
     "等交互功能则以页面走查验证，全部用例与结果见下表；页面走查覆盖学生端 9 页、管理端 10 页在桌面/iOS/Android、中英双语、"
     "明暗主题下的表现，均无控制台错误。",
     []),
    ("3.2  登录与权限",
     "系统提供学生/管理员两类角色登录及注册、找回密码、图形验证码等能力，登录后由 Sa-Token 维护登录态与角色鉴权，"
     "并按角色路由到对应工作台。登录页如下图所示。",
     [(sp("scripts/report/desktop/00-login.png"), "登录页")]),
    ("3.3  选座预约主流程",
     "学生按“校区→楼栋→楼层”筛选自习室（下图），进入后在座位网格上选择日期、时间片与具体座位提交预约；"
     "座位状态（空闲/已约/使用中/禁用）以颜色区分，最终预约结果以后端返回为准。",
     [(sp("scripts/report/desktop/student-rooms.png"), "自习室筛选与列表"),
      (sp("scripts/report/shots/03-student-seats.png"), "座位网格选座")]),
    ("3.4  实时热力看板",
     "管理端与学生端共享“快照 + SSE”实时看板：两个客户端打开同一自习室，一端预约/释放，另一端座位秒级变色；"
     "时空占用图进一步以时间轴维度展示各时段占用情况，验证了 SSE 增量推送与多端一致性。",
     [(sp("scripts/report/shots/09-admin-board.png"), "管理端实时看板"),
      (sp("scripts/report/desktop/admin-spacetime.png"), "时空占用图")]),
    ("3.5  管理与运营",
     "管理端提供概览驾驶舱、座位排布可视化编辑器与黑名单管理等能力，支持自习室/座位维护、爽约用户查看与解禁等运营操作。"
     "座位排布编辑器支持逐格设置座位/过道/空位/禁用位以适配多过道与不规则房间，并可在网格四周放置“门”与“讲台”边缘标记"
     "（不占用座位），使座位图与真实教室方位一致，便于学生按门与讲台的位置选座。",
     [(sp("scripts/report/desktop/admin-dashboard.png"), "管理台概览"),
      (sp("scripts/report/shots/42-门和讲台的形式.png"), "座位排布编辑器（含门 / 讲台边缘标记）"),
      (sp("scripts/report/shots/11-admin-blacklist.png"), "黑名单管理")]),
    ("3.6  数据报表与积分排行",
     "系统基于聚合口径生成使用率、取消率、爽约率、热门时段与利用率排行等报表，并按守约/爽约规则结算积分、生成积分排行榜。",
     [(sp("scripts/report/desktop/admin-reports.png"), "数据报表"),
      (sp("scripts/report/desktop/student-ranking.png"), "积分排行榜")]),
    ("3.7  地图选点与定位推荐最近自习室",
     "为支撑“按距离推荐”，系统接入地图服务（Leaflet + 高德地图瓦片）实现楼栋坐标的可视化维护：管理员在“位置管理”中点击"
     "「地图选点」，即可在卫星底图上直接点选或拖拽标记确定楼栋位置，实时回显所选经纬度（下图中为“图书馆C座”选点，"
     "纬度 24.605766、经度 118.311826），确认后保存，免去手工输入坐标的麻烦。"
     "学生端“附近空位”则据此实现定位推荐：点击「定位」后读取浏览器当前位置，与各楼栋坐标计算距离，自动识别最近楼栋"
     "（下图中定位到“图书馆A座”，约 1742 m），并按“同楼栋 > 距离最近 > 空位更多”的顺序推荐仍有空位的自习室，"
     "卡片上直接给出距离与实时空位数；无空位的自习室不会被返回。两者构成“管理员维护坐标 → 学生按距离找座”的完整闭环。",
     [(sp("scripts/report/shots/45-接入地图.png"), "位置管理 · 地图选点维护楼栋坐标"),
      (sp("scripts/report/shots/43-定位.png"), "定位推荐最近有空位的自习室")]),
    ("3.8  专注番茄钟",
     "学生端内置专注番茄钟，采用番茄工作法（25 分钟专注 + 5 分钟短休，每 4 个番茄享受一次长休），"
     "支持开始/暂停、重置、跳过与三种模式切换，并统计今日完成番茄数、本轮进度与今日专注时长，"
     "与自习预约配合帮助学生在自习室内保持专注、量化专注成果。",
     [(sp("scripts/report/shots/44-番茄钟.png"), "专注番茄钟（专注计时进行中）")]),
    ("3.9  智能与工程化",
     "学生端集成 AI 智能选座助手，可用自然语言表达需求并给出可解释的座位推荐（离线内置规则引擎，亦可接入大模型）；"
     "后端接口经 Knife4j 生成可交互文档，便于联调与验收。",
     [(sp("scripts/report/shots/13-ai-assistant.png"), "AI 选座助手"),
      (sp("scripts/report/shots/12-knife4j.png"), "Knife4j 接口文档")]),
    ("3.10  多端适配与国际化",
     "前端实现了明暗主题、移动端响应式（iOS / Android）与中英双语国际化。在 iPhone 13 与 Pixel 5 模拟设备上，"
     "侧边栏收起为抽屉式导航、栅格纵向堆叠，无横向溢出；语言切换即时联动界面与组件。",
     [(sp("scripts/report/shots/27-dark-student.png"), "深色模式"),
      (sp("scripts/report/ios/student-home.png"), "iOS 移动端"),
      (sp("scripts/report/android/admin-reports.png"), "Android 移动端"),
      (sp("scripts/report/desktop-en/student-home.png"), "英文界面")]),
]

TEST_TABLE = (
    ["编号", "测试项", "操作 / 输入", "预期结果", "结果"],
    [
        ["T01", "登录鉴权", "演示账号登录", "返回 token 与角色，路由正确", "通过"],
        ["T02", "自习室查询", "按校区/楼栋/楼层筛选", "返回匹配自习室与开放状态", "通过"],
        ["T03", "看板快照", "打开自习室看板", "返回座位网格与每格状态", "通过"],
        ["T04", "正常预约", "选座+时间片提交", "预约成功并占用时间片", "通过"],
        ["T05", "重复预约拒绝", "同座同片再次预约", "返回 SEAT_ALREADY_RESERVED", "通过"],
        ["T06", "并发抢座", "8 人并发抢同一座位", "仅 1 人成功，其余被拒", "通过"],
        ["T07", "签到 / 签退", "预约后签到再签退", "状态流转正确并释放座位", "通过"],
        ["T08", "超时释放", "超签到窗口未签到", "自动释放并计爽约", "通过"],
        ["T09", "单日限次", "超过单日预约次数", "返回 DAILY_LIMIT_EXCEEDED", "通过"],
        ["T10", "黑名单", "爽约达阈值后预约", "返回 USER_IN_BLACKLIST", "通过"],
        ["T11", "实时推送", "一端预约另一端观察", "另一端座位秒级变色", "通过"],
        ["T12", "数据报表", "查看报表页", "口径与聚合一致", "通过"],
        ["T13", "权限校验", "学生访问管理接口", "被拒绝（无权限）", "通过"],
        ["T14", "地图选点", "位置管理中地图点选楼栋坐标", "回显经纬度并保存成功", "通过"],
        ["T15", "定位推荐", "点击定位查找附近空位", "识别最近楼栋并按距离+空位推荐", "通过"],
        ["T16", "专注番茄钟", "开始 25 分钟专注计时", "计时运行并统计番茄数/时长", "通过"],
    ],
    "主要功能测试用例与结果",
)


CODE_SLOT = """// slotIndex = (hour*60+minute)/slotMinutes，从 00:00 起（30 分钟片时 14:00->28）
public static int toSlot(LocalTime time, int slotMinutes) {
    return (time.getHour() * 60 + time.getMinute()) / slotMinutes;
}
// 将预约区间 [startSlot, endSlot) 展开为占用的时间片序号列表
public static List<Integer> expand(int startSlot, int endSlot) {
    List<Integer> list = new ArrayList<>();
    for (int i = startSlot; i < endSlot; i++) list.add(i);
    return list;
}"""

CODE_RESERVE = """// Redisson 加锁降低冲突（锁失败或失效仍由数据库唯一索引兜底）
RLock lock = redisson.getLock("seat:"+seatId+":date:"+date+":slots:"+startSlot+"-"+endSlot);
boolean locked = lock.tryLock(3, 10, TimeUnit.SECONDS);
try {
    reservation = tx.execute(status -> {          // 同一事务内写入
        Reservation r = new Reservation(userId, seatId, roomId, date, startSlot, endSlot);
        r.setStatus("PENDING_SIGN_IN");
        reservationMapper.insert(r);
        for (Integer s : slots) {                 // 逐片插入占用记录
            slotMapper.insert(new ReservationSlot(r.getId(), seatId, date, s));
        }                                         // 命中 uk(seat_id,date,slot_index) 即冲突
        return r;
    });
} catch (DuplicateKeyException e) {               // 并发双占被数据库原子拒绝
    throw new BizException(BizError.SEAT_ALREADY_RESERVED);
} finally {
    if (locked && lock.isHeldByCurrentThread()) lock.unlock();
}"""

CODE_SSE = """// 座位状态变更后，向所有订阅该看板的客户端广播增量事件
private void broadcastSeat(Long roomId, LocalDate date, Long seatId,
                          String status, String event) {
    Map<String, Object> payload = new HashMap<>();
    payload.put("seatId", seatId);
    payload.put("status", status);        // FREE / RESERVED / USING
    sse.broadcast(roomId, date, event, payload);   // event 如 seat_reserved
}"""


def build_design(c):
    c.body("本部分从总体架构、数据库、并发模型、核心流程、状态机、实时看板、超时释放、部署与接口九个方面阐述系统设计。"
           "系统的两大技术核心是“时间片并发选座的正确性”与“多客户端实时看板的一致性”，其余设计均围绕二者展开。")
    for title, prose, img, name in DESIGN_FIGS:
        c.sub(title)
        c.body(prose)
        c.figure(img, name)
        if title.split()[0] == "2.2":
            c.table(DB_TABLE[0], DB_TABLE[1], DB_TABLE[2], col_w=[3.2, 2.4, 1.6, 5.0, 3.2])
        elif title.split()[0] == "2.3":
            c.body("时间片换算与拆分的核心代码如下：slotIndex 由“(时*60+分)/片长”计算，预约区间按 [起, 止) 展开为片序号列表。")
            c.code("时间片换算与拆分（SlotUtil）", CODE_SLOT)
        elif title.split()[0] == "2.4":
            c.body("并发预约的核心代码如下：先用 Redisson 加锁降低冲突，再在同一事务内插入主记录与逐片占用记录；"
                   "一旦命中唯一索引即抛出 DuplicateKeyException 并回滚，返回“座位已被预约”，从而保证绝不双占。")
            c.code("并发预约与唯一索引兜底（ReservationService）", CODE_RESERVE)
        elif title.split()[0] == "2.6":
            c.body("座位状态变更后广播 SSE 增量事件的核心代码如下：")
            c.code("SSE 增量广播（ReservationService）", CODE_SSE)
    c.sub("2.9  核心接口概览")
    c.body("前后端通过 REST 接口与 SSE 通道交互，统一响应结构为 code/message/data/traceId。核心接口一览见下表。")
    c.table(API_TABLE[0], API_TABLE[1], API_TABLE[2], col_w=[2.4, 2.0, 6.2, 4.8])


def build_test(c):
    c.body("本部分给出系统的测试环境、方法与主要功能的测试结果。测试以自动化脚本验证核心正确性，"
           "以真实浏览器多端走查验证交互与展示效果。")
    for title, prose, imgs in TEST_SECTIONS:
        c.sub(title)
        c.body(prose)
        for img_abs, name in imgs:
            c.figure(img_abs, name)
        if title.split()[0] == "3.1":
            c.table(TEST_TABLE[0], TEST_TABLE[1], TEST_TABLE[2], col_w=[1.4, 2.6, 4.4, 4.6, 1.4])
    c.sub("3.11  测试结论")
    c.body("经自动化脚本与多端页面走查验证，系统在并发正确性、实时一致性、座位生命周期闭合、权限控制等核心红线上均达标，"
           "13 项核心用例全部通过，桌面与移动多端、中英双语、明暗主题下均运行正常、无控制台错误，满足预期设计目标。")


# ============================================================ 组装
def build(images_only):
    _fig[0] = 0; _tbl[0] = 0
    doc = docx.Document(TEMPLATE)
    H_photo = find_p(doc, "项目组合影和分工介绍")
    H_goal = find_p(doc, "项目目标与要求")
    H_content = find_p(doc, "项目内容")
    H_func = find_p(doc, "（一）、功能需求")
    H_design = find_p(doc, "（二）、系统设计")
    H_test = find_p(doc, "（三）、系统测试")
    for name, el in [("photo", H_photo), ("goal", H_goal), ("content", H_content),
                     ("func", H_func), ("design", H_design), ("test", H_test)]:
        assert el is not None, f"未找到标题：{name}"

    # 封面：项目组名 / 选题 / 成员
    fill_title(doc)

    # 清理模板占位
    delete_between(H_photo, H_goal)
    delete_between(H_goal, H_content)
    delete_between(H_content, H_func)
    delete_between(H_func, H_design)
    delete_between(H_design, H_test)
    delete_after_until_sectpr(H_test)

    # 插入正文
    build_photo(Cursor(doc, H_photo, images_only))
    build_goal(Cursor(doc, H_goal, images_only))
    build_content_intro(Cursor(doc, H_content, images_only))
    build_func(Cursor(doc, H_func, images_only))
    build_design(Cursor(doc, H_design, images_only))
    build_test(Cursor(doc, H_test, images_only))

    out = OUT_IMG if images_only else OUT_FULL
    doc.save(out)
    print("已生成:", out)


if __name__ == "__main__":
    build(images_only=False)  # 仅生成完整版
    print("完成。")
